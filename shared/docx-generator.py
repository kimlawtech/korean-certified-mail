#!/usr/bin/env python3
"""
korean-contracts DOCX 생성기
사용법: python3 docx-generator.py <input.md> [output.docx]
"""

import sys
import re
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("python-docx 설치 필요: pip3 install python-docx")
    sys.exit(1)


FONT_NAME = "굴림"
FONT_NAME_EN = "Arial"

# 스타일 정의
STYLES = {
    "disclaimer": {"size": 9,  "color": RGBColor(0x88, 0x88, 0x88), "italic": True},
    "h1":         {"size": 16, "color": RGBColor(0x1A, 0x1A, 0x2E), "bold": True},
    "h2":         {"size": 13, "color": RGBColor(0x16, 0x21, 0x3E), "bold": True},
    "h3":         {"size": 11, "color": RGBColor(0x0F, 0x3A, 0x57), "bold": True},
    "body":       {"size": 10, "color": RGBColor(0x1A, 0x1A, 0x1A), "bold": False},
    "warning":    {"size": 9,  "color": RGBColor(0xCC, 0x44, 0x00), "bold": True},
    "bold_body":  {"size": 10, "color": RGBColor(0x1A, 0x1A, 0x1A), "bold": True},
}


def set_font(run, style_key):
    s = STYLES[style_key]
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(s["size"])
    run.font.color.rgb = s["color"]
    run.font.bold = s.get("bold", False)
    run.font.italic = s.get("italic", False)


def set_para_font(para, style_key):
    for run in para.runs:
        set_font(run, style_key)


def add_paragraph(doc, text, style_key="body", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = Pt(STYLES[style_key]["size"] * 1.6)

    # 굵기 마크다운 파싱 (**text**)
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        s = STYLES[style_key]
        run.font.size = Pt(s["size"])
        run.font.color.rgb = s["color"]
        run.font.bold = s.get("bold", False) or (i % 2 == 1)
        run.font.italic = s.get("italic", False)
    return para


def add_divider(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_table(doc, rows):
    """rows: list of list of str (first row = header if detected)"""
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j >= col_count:
                break
            cell = row.cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            run = para.add_run(cell_text.strip())
            run.font.name = FONT_NAME
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            if i == 0:
                run.font.bold = True
                shading = OxmlElement("w:shd")
                shading.set(qn("w:val"), "clear")
                shading.set(qn("w:color"), "auto")
                shading.set(qn("w:fill"), "E8EDF3")
                cell._tc.get_or_add_tcPr().append(shading)

    # 컬럼 너비 자동 균등
    table_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
    col_width = table_width // col_count
    for col in table.columns:
        for cell in col.cells:
            cell.width = col_width

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def page_setup(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)


def parse_and_render(md_text, doc):
    lines = md_text.split("\n")
    table_rows = []
    in_table = False

    i = 0
    while i < len(lines):
        line = lines[i]

        # 테이블 처리
        if line.strip().startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            # 구분선 행 스킵
            if all(re.match(r"^-+$", c) for c in cells if c):
                i += 1
                continue
            table_rows.append(cells)
            in_table = True
            i += 1
            continue
        else:
            if in_table:
                add_table(doc, table_rows)
                table_rows = []
                in_table = False

        stripped = line.strip()

        # 빈 줄
        if not stripped:
            i += 1
            continue

        # 면책 blockquote
        if stripped.startswith("> "):
            text = stripped[2:]
            add_paragraph(doc, text, "disclaimer", space_before=0, space_after=2)
            i += 1
            continue

        # 경고 (⚠️ 또는 ※)
        if stripped.startswith("⚠️") or stripped.startswith("※"):
            add_paragraph(doc, stripped, "warning", space_before=4, space_after=4)
            i += 1
            continue

        # 구분선
        if stripped == "---":
            add_divider(doc)
            i += 1
            continue

        # 제목
        if stripped.startswith("### "):
            p = add_paragraph(doc, stripped[4:], "h3", space_before=8, space_after=3)
            i += 1
            continue
        if stripped.startswith("## "):
            p = add_paragraph(doc, stripped[3:], "h2", space_before=12, space_after=4)
            i += 1
            continue
        if stripped.startswith("# "):
            p = add_paragraph(doc, stripped[2:], "h1", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=10)
            i += 1
            continue

        # 본문
        add_paragraph(doc, stripped, "body", space_before=0, space_after=3)
        i += 1

    # 마지막 테이블 flush
    if in_table and table_rows:
        add_table(doc, table_rows)


def generate_docx(md_path: str, out_path: str = None):
    md_text = Path(md_path).read_text(encoding="utf-8")

    if not out_path:
        stem = Path(md_path).stem
        out_path = str(Path(md_path).parent / f"{stem}.docx")

    doc = Document()
    page_setup(doc)

    # 기본 스타일 폰트 설정
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    parse_and_render(md_text, doc)
    doc.save(out_path)
    print(f"✓ 저장됨: {out_path}")
    return out_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("사용법: python3 docx-generator.py <input.md> [output.docx]")
        sys.exit(1)
    md_path = sys.argv[1]
    out_path = sys.argv[2] if len(sys.argv) > 2 else None
    generate_docx(md_path, out_path)
