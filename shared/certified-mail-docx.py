#!/usr/bin/env python3
"""
내용증명 전용 DOCX 생성기

한국 우체국·법원 제출용 내용증명 서식 기준:
- 폰트: 바탕체 (한글 본문 표준), 제목 HY헤드라인M
- 용지: A4, 좌3cm 우2.5cm 상하2.5cm
- 줄간격: 200% (법원 제출 기준)
- 제목 가운데 정렬, 본문 양쪽 정렬
- 당사자 정보 표 형식
- 구분선 실선

사용법:
  python3 certified-mail-docx.py <input.txt> [output.docx]

TXT 파일 구조 파싱:
  TITLE: 제목 한 줄
  SENDER: 발신인 블록 (빈 줄로 구분)
  RECIPIENT: 수신인 블록
  DATE: 작성일
  BODY: 본문 (번호 조항 포함)
  SIGNATURE: 서명란
"""

import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("python-docx 설치 필요: pip3 install python-docx")
    sys.exit(1)


# ── 폰트 설정 ─────────────────────────────────────────────
FONT_BODY   = "바탕"          # 본문: 바탕체 (법원·공문서 표준)
FONT_TITLE  = "바탕"          # 제목: 바탕체 (굵게 처리)
FONT_LATIN  = "Times New Roman"  # 영문·숫자 (법원 표준)

# ── 색상 ──────────────────────────────────────────────────
COLOR_BLACK      = RGBColor(0x00, 0x00, 0x00)
COLOR_DISCLAIMER = RGBColor(0x55, 0x55, 0x55)


# ══════════════════════════════════════════════════════════
# 페이지 설정
# ══════════════════════════════════════════════════════════

def page_setup(doc):
    """A4, 여백: 좌3cm 우2.5cm 상하2.5cm (법원 제출 기준)"""
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)


# ══════════════════════════════════════════════════════════
# 폰트·스타일 헬퍼
# ══════════════════════════════════════════════════════════

def _apply_font(run, size_pt, bold=False, italic=False, color=COLOR_BLACK, font_name=FONT_BODY):
    run.font.name  = FONT_LATIN
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.color.rgb = color
    # 동아시아 폰트 직접 지정 (한글 폰트 보장)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:hAnsi"),    FONT_LATIN)
    rFonts.set(qn("w:ascii"),    FONT_LATIN)


def _set_line_spacing(para, spacing_pct=200):
    """줄간격 설정 (200% = 법원 제출 기준)"""
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    # 200% = 240 twips × 2 = 480; WD_LINE_SPACING.MULTIPLE + 규칙
    spacing.set(qn("w:line"),     str(int(240 * spacing_pct / 100)))
    spacing.set(qn("w:lineRule"), "auto")
    existing = pPr.find(qn("w:spacing"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(spacing)


def _add_border_bottom(para, color="000000", size=12):
    """단락 아래 실선 테두리"""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(size))   # 1/8 pt 단위
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


# ══════════════════════════════════════════════════════════
# 단락 추가 헬퍼
# ══════════════════════════════════════════════════════════

def add_para(doc, text, size=11, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, color=COLOR_BLACK,
             space_before=0, space_after=6, line_spacing=200,
             font_name=FONT_BODY):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    _set_line_spacing(para, line_spacing)

    if text:
        run = para.add_run(text)
        _apply_font(run, size, bold=bold, italic=italic, color=color, font_name=font_name)
    return para


def add_empty(doc, space_after=6):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(space_after)
    return para


# ══════════════════════════════════════════════════════════
# 구성 요소
# ══════════════════════════════════════════════════════════

def add_disclaimer(doc, text):
    """면책 문구 — 회색 이탤릭, 작은 글씨"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(4)
    _set_line_spacing(para, 160)
    run = para.add_run(text)
    _apply_font(run, 8.5, italic=True, color=COLOR_DISCLAIMER)
    return para


def add_thick_rule(doc):
    """굵은 실선 (제목 위아래)"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "bottom"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "16")   # 2pt
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), "000000")
        pBdr.append(el)
    pPr.append(pBdr)
    return para


def add_thin_rule(doc):
    """얇은 실선 (섹션 구분)"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    _add_border_bottom(para, color="888888", size=4)
    return para


def add_title_block(doc, title_text):
    """
    내용증명 제목 블록
    ═══════════════════
    내  용  증  명  서
    ═══════════════════
    """
    add_thick_rule(doc)
    # 제목: 16pt 굵게 가운데 정렬, 자간 넓게
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(6)
    _set_line_spacing(para, 150)
    run = para.add_run(title_text)
    _apply_font(run, 16, bold=True, font_name=FONT_TITLE)
    # 자간 넓히기 (200 = 20pt × 1/100pt 단위)
    rPr = run._element.get_or_add_rPr()
    spacing_el = OxmlElement("w:spacing")
    spacing_el.set(qn("w:val"), "200")
    rPr.append(spacing_el)
    add_thick_rule(doc)
    return para


def add_parties_table(doc, parties: list[dict]):
    """
    당사자 정보 표
    parties = [
        {"label": "발 신 인", "lines": ["홍길동", "서울시 ...", "010-..."]},
        {"label": "수 신 인", "lines": ["이임대인", "서울시 ..."]},
    ]
    """
    rows = len(parties)
    table = doc.add_table(rows=rows, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 컬럼 너비: 라벨 3.5cm, 내용 나머지
    page_w = Cm(21.0) - Cm(3.0) - Cm(2.5)  # 15.5cm
    col_label_w = Cm(3.5)
    col_val_w   = page_w - col_label_w

    for i, party in enumerate(parties):
        row = table.rows[i]

        # 라벨 셀
        label_cell = row.cells[0]
        label_cell.width = col_label_w
        _set_cell_shading(label_cell, "F2F2F2")
        para = label_cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after  = Pt(4)
        _set_line_spacing(para, 160)
        run = para.add_run(party["label"])
        _apply_font(run, 10, bold=True)

        # 내용 셀
        val_cell = row.cells[1]
        val_cell.width = col_val_w
        for j, line in enumerate(party["lines"]):
            if j == 0:
                para = val_cell.paragraphs[0]
            else:
                para = val_cell.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            _set_line_spacing(para, 160)
            run = para.add_run(line)
            _apply_font(run, 10)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _set_cell_shading(cell, fill_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_pr.append(shd)


def add_subject_line(doc, subject):
    """제 목: XXX"""
    add_thin_rule(doc)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(6)
    _set_line_spacing(para, 160)
    run_label = para.add_run("제    목:  ")
    _apply_font(run_label, 11, bold=True)
    run_val = para.add_run(subject)
    _apply_font(run_val, 11, bold=True)
    add_thin_rule(doc)


def add_body_section(doc, number, title, content_lines):
    """
    번호 조항 추가
    예) 1. 임대차계약 내용
        내용...
    """
    add_empty(doc, space_after=4)
    # 조항 제목
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(2)
    _set_line_spacing(para, 200)
    heading = f"{number}. {title}" if number else title
    run = para.add_run(heading)
    _apply_font(run, 11, bold=True)

    # 조항 내용
    for line in content_lines:
        if not line.strip():
            add_empty(doc, space_after=2)
            continue
        # 소항목 (가. 나. 다. 또는 ① ② 또는 - 또는 · 로 시작)
        is_sub = bool(re.match(r"^\s*([가-힣]\.|①|②|③|④|⑤|-|·|\*)", line.strip()))
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(3)
        _set_line_spacing(para, 200)
        if is_sub:
            para.paragraph_format.left_indent = Cm(0.8)
        run = para.add_run(line.strip())
        _apply_font(run, 11)


def add_signature_block(doc, date_str, sender_name, account_info=None):
    """
    서명란
    """
    add_empty(doc, space_after=10)
    add_thin_rule(doc)

    # 작성일
    para = add_para(doc, date_str, size=11, bold=False,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before=8, space_after=6)

    # 계좌 정보 (있으면)
    if account_info:
        add_para(doc, f"입금 계좌: {account_info}", size=10,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=8)

    # 서명
    add_empty(doc, space_after=16)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    _set_line_spacing(para, 160)
    run = para.add_run(f"발신인:  {sender_name}  (서명 또는 날인)")
    _apply_font(run, 11)

    add_empty(doc, space_after=4)
    add_thin_rule(doc)


def add_dispatch_guide(doc):
    """우체국 발송 안내"""
    add_empty(doc, space_after=8)
    guide_title = "[ 우체국 내용증명 발송 방법 ]"
    add_para(doc, guide_title, size=9, bold=True,
             align=WD_ALIGN_PARAGRAPH.LEFT,
             color=RGBColor(0x33, 0x33, 0x33), space_before=4, space_after=2,
             line_spacing=160)
    lines = [
        "1. 동일 내용의 문서 3부 출력 (발신인 보관 1부 / 수신인 발송 1부 / 우체국 보관 1부)",
        "2. 가까운 우체국 방문 또는 인터넷우체국(epost.go.kr) → e-내용증명 이용 가능",
        "3. 발송 후 등기 접수증을 반드시 보관하십시오 (향후 소송 시 증거 서류)",
        "※ 내용증명 자체에 법적 강제력은 없으며, 발송 사실과 내용의 공적 증명이 효력입니다.",
        "※ 미이행 시 임차권등기명령·소액심판·민사소송 등 후속 절차를 진행하십시오.",
    ]
    for line in lines:
        add_para(doc, line, size=8.5, italic=False,
                 color=RGBColor(0x44, 0x44, 0x44),
                 align=WD_ALIGN_PARAGRAPH.LEFT,
                 space_before=0, space_after=1, line_spacing=160)


# ══════════════════════════════════════════════════════════
# TXT 파싱 → 구조 추출
# ══════════════════════════════════════════════════════════

def parse_certified_mail_txt(text: str) -> dict:
    """
    내용증명 TXT를 구조화된 dict로 파싱한다.

    지원 헤더 태그:
    [TITLE], [SENDER], [RECIPIENT], [DATE], [SUBJECT], [SECTION_N], [SIGNATURE], [ACCOUNT]

    태그 없는 경우: 직접 텍스트 파싱으로 폴백
    """
    result = {
        "title":      "내  용  증  명  서",
        "sender":     [],
        "recipient":  [],
        "date":       "",
        "subject":    "",
        "sections":   [],   # [{number, title, lines}]
        "signature":  "",
        "account":    "",
        "disclaimer": "",
    }

    lines = text.split("\n")
    current_tag = None
    current_buf = []

    def flush(tag, buf):
        content = "\n".join(buf).strip()
        if not content:
            return
        if tag == "TITLE":
            result["title"] = content
        elif tag == "SENDER":
            result["sender"] = [l.strip() for l in buf if l.strip()]
        elif tag == "RECIPIENT":
            result["recipient"] = [l.strip() for l in buf if l.strip()]
        elif tag == "DATE":
            result["date"] = content
        elif tag == "SUBJECT":
            result["subject"] = content
        elif tag == "SIGNATURE":
            result["signature"] = content
        elif tag == "ACCOUNT":
            result["account"] = content
        elif tag and tag.startswith("SECTION"):
            parts = content.split("\n", 1)
            sec_title = parts[0].strip()
            sec_lines = parts[1].split("\n") if len(parts) > 1 else []
            # 번호 자동 추출 (예: "1. 제목" → number=1, title="제목")
            m = re.match(r"^(\d+)\.\s+(.+)$", sec_title)
            if m:
                result["sections"].append({
                    "number": m.group(1),
                    "title":  m.group(2),
                    "lines":  sec_lines,
                })
            else:
                result["sections"].append({
                    "number": "",
                    "title":  sec_title,
                    "lines":  sec_lines,
                })
        elif tag == "DISCLAIMER":
            result["disclaimer"] = content

    # 태그 기반 파싱 먼저 시도
    tagged = False
    for line in lines:
        m = re.match(r"^\[([A-Z_0-9]+)\]\s*$", line.strip())
        if m:
            tagged = True
            flush(current_tag, current_buf)
            current_tag = m.group(1)
            current_buf = []
        else:
            current_buf.append(line)
    flush(current_tag, current_buf)

    if tagged:
        return result

    # ── 태그 없는 경우: 텍스트 휴리스틱 파싱 ──────────────────
    # 면책 문구 (첫 번째 ※ 또는 > 로 시작하는 블록)
    i = 0
    while i < len(lines) and (lines[i].startswith("※") or lines[i].startswith(">")):
        result["disclaimer"] += lines[i].lstrip(">").strip() + "\n"
        i += 1

    # 발신인/수신인 블록 탐지
    in_parties = False
    parties_done = False
    body_lines = []

    for line in lines[i:]:
        stripped = line.strip()

        if not parties_done:
            # "발 신 인" 또는 "발신인" 패턴
            m_sender = re.match(r"^발\s*신\s*인\s*[:：]?\s*(.*)$", stripped)
            m_recip  = re.match(r"^수\s*신\s*인\s*[:：]?\s*(.*)$", stripped)
            m_date   = re.match(r"^(작성일|작 성 일)\s*[:：]?\s*(.+)$", stripped)
            m_subj   = re.match(r"^(제\s*목|제    목)\s*[:：]?\s*(.+)$", stripped)

            if m_sender:
                if m_sender.group(1):
                    result["sender"].append(m_sender.group(1).strip())
                in_parties = "sender"
                continue
            if m_recip:
                if m_recip.group(1):
                    result["recipient"].append(m_recip.group(1).strip())
                in_parties = "recipient"
                continue
            if m_date:
                result["date"] = m_date.group(2).strip()
                parties_done = True
                continue
            if m_subj:
                result["subject"] = m_subj.group(2).strip()
                continue

            if in_parties == "sender" and stripped and not re.match(r"^수\s*신\s*인", stripped):
                if not re.match(r"^[-━─=]+$", stripped):
                    result["sender"].append(stripped)
                continue
            if in_parties == "recipient" and stripped:
                if not re.match(r"^[-━─=]+$", stripped) and not re.match(r"^(제|작성일)", stripped):
                    result["recipient"].append(stripped)
                continue

        else:
            body_lines.append(line)

    # 본문 섹션 분리 (숫자. 제목 패턴으로 분리)
    current_sec = None
    for line in body_lines:
        stripped = line.strip()
        m_sec = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if m_sec and len(stripped) < 50:  # 짧으면 제목으로 판단
            if current_sec:
                result["sections"].append(current_sec)
            current_sec = {
                "number": m_sec.group(1),
                "title":  m_sec.group(2),
                "lines":  [],
            }
        elif current_sec is not None:
            current_sec["lines"].append(line)
        elif stripped and not re.match(r"^[-━─=]+$", stripped):
            # 섹션 외 본문: 섹션 0으로 처리
            if not result["sections"] or result["sections"][-1]["number"] != "":
                result["sections"].append({"number": "", "title": "", "lines": [line]})
            else:
                result["sections"][-1]["lines"].append(line)

    if current_sec:
        result["sections"].append(current_sec)

    return result


# ══════════════════════════════════════════════════════════
# 메인 렌더러
# ══════════════════════════════════════════════════════════

def render_certified_mail(parsed: dict, out_path: str):
    doc = Document()
    page_setup(doc)

    # 기본 스타일 폰트 설정
    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal.font.size = Pt(11)
    nf = normal._element.get_or_add_rPr().find(qn("w:rFonts"))
    if nf is None:
        nf = OxmlElement("w:rFonts")
        normal._element.get_or_add_rPr().insert(0, nf)
    nf.set(qn("w:eastAsia"), FONT_BODY)

    # ① 면책 문구
    if parsed.get("disclaimer"):
        add_disclaimer(doc, parsed["disclaimer"].strip())
        add_empty(doc, space_after=8)

    # ② 제목 블록
    add_title_block(doc, parsed["title"])
    add_empty(doc, space_after=6)

    # ③ 당사자 표
    parties = []
    if parsed["sender"]:
        # 라벨 자간 맞추기
        parties.append({"label": "발  신  인", "lines": parsed["sender"]})
    if parsed["recipient"]:
        parties.append({"label": "수  신  인", "lines": parsed["recipient"]})
    if parties:
        add_parties_table(doc, parties)

    # ④ 제목 줄
    if parsed.get("subject"):
        add_subject_line(doc, parsed["subject"])
        add_empty(doc, space_after=8)

    # ⑤ 본문 섹션
    for sec in parsed["sections"]:
        add_body_section(doc, sec["number"], sec["title"], sec["lines"])

    # ⑥ 서명란
    signature_name = parsed.get("signature") or (parsed["sender"][0] if parsed["sender"] else "")
    date_str = parsed.get("date") or ""
    account  = parsed.get("account") or ""
    if date_str or signature_name:
        add_signature_block(doc, date_str, signature_name, account_info=account or None)

    # ⑦ 발송 안내
    add_dispatch_guide(doc)

    doc.save(out_path)
    print(f"저장됨: {out_path}")
    return out_path


# ══════════════════════════════════════════════════════════
# 진입점
# ══════════════════════════════════════════════════════════

def generate(input_path: str, out_path: str = None):
    text = Path(input_path).read_text(encoding="utf-8")
    if not out_path:
        stem = Path(input_path).stem
        out_path = str(Path(input_path).parent / f"{stem}.docx")
    parsed = parse_certified_mail_txt(text)
    return render_certified_mail(parsed, out_path)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("사용법: python3 certified-mail-docx.py <input.txt> [output.docx]")
        sys.exit(1)
    generate(sys.argv[1], sys.argv[2] if len(sys.argv) > 2 else None)
